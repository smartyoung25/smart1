# -*- coding: utf-8 -*-
"""KAASA Farmingsight — 7단계 실증 추진 로드맵 (Word .docx)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN_DARK = RGBColor(0x0F, 0x51, 0x32)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
INK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x60, 0x5A)
FONT = "맑은 고딕"

STEPS = [
    (1, "참여농가 사전진단 및 현장 문제 유형화",
     "C17 시스템 종합진단 · C18 현장 컨설팅 문진",
     [("C17 시스템 종합진단",
       "시스템 성숙도 종합점수와 진단 추세(회차별), 6대 영역 진단, 우선순위 처방(의사결정)·ROI, "
       "인프라 성숙도 상세·업그레이드 추천, 결과보고서 PDF를 제공한다."),
      ("C18 현장 컨설팅 문진",
       "RDA 현장 컨설팅 체크리스트로 원수·필터·유황훈증·병해충·인증 등 센서가 보지 못하는 항목까지 "
       "상태(불량·주의·메모)를 입력해 종합진단에 반영하여 현장 문제를 유형화한다.")]),
    (2, "데이터 수집항목·수집주기·입력방식 확정",
     "C21 연동·서비스 신청 · C2 데이터 동의",
     [("C21 연동·서비스 신청",
       "장비/서비스명·제조사·통신 프로토콜·현재 인프라 정도(①미구축 ~ ③원격 가능)·연락처·요청내용으로 "
       "이기종 장비·외부 데이터·전문가 연동을 신청하고 신청 내역을 관리한다."),
      ("C2 데이터 동의",
       "데이터 활용 모드와 항목별 공개 토글(민감 경영정보 비공개, 학습·벤치마킹에만 활용)로 "
       "수집항목·입력방식·활용범위를 확정한다.")]),
    (3, "농가별 데이터 수집 및 품질관리",
     "C16 시설기자재 등록 · C8 이기종 연동",
     [("C16 시설기자재 등록",
       "제조사·프로토콜이 다른 장비를 등록하고 데이터포인트를 표준변수로 매핑한다"
       "(견적서·시방서·엑셀·PDF 자동추출 후 검토 등록). 분류·장비ID·주소·포트·통신 프로토콜을 입력한다."),
      ("C8 이기종 연동",
       "연동 API 수·연결됨·연동률, 외부 데이터 연동현황(실시간), MQTT·게이트웨이, 업그레이드 추천으로 "
       "수집 데이터의 품질·연동 상태를 관리한다.")]),
    (4, "AI 의사결정 지원 대시보드 시범적용",
     "C3 통합 홈 · G2 환경·기후 제어",
     [("C3 통합 홈",
       "오늘의 결정(AI 의사결정)·수확예측·예상매출·배액률·VPD·우수농가 벤치마크·오늘 할 일·"
       "AI 추천(관수 신뢰도)·온실/노지/출하 현황 요약을 통합 표시한다."),
      ("G2 환경·기후 제어",
       "온도·습도·CO₂·VPD·DLI 실측, 환경 추이, 제어모드 4단계(수동보조·권고형·승인형·자동), "
       "환경관리 전략표, AI 에이전트(30초 루프), 외기·에너지요금으로 AI 의사결정을 시범적용한다.")]),
    (5, "농가별 분석리포트 제공 및 현장 컨설팅",
     "C14 월간 경영성과 리포트 · C4 AI 진단",
     [("C14 월간 경영성과 리포트",
       "6대 영역 종합·전월 대비 변화·9대 성과지표(목표 대비)·데이터 학습 기여·이행활동·"
       "현장점검 진단·취약항목·이달의 실행 To-do·PDF 내보내기를 제공한다."),
      ("C4 AI 진단",
       "AI 종합진단 점수, AI 1차 개선추천, 환경이상 감지, 전문가 컨설팅 예약"
       "(재배기술·병해방제·경영 컨설턴트, 현장점검 결과 자동첨부)으로 리포트·컨설팅을 제공한다.")]),
    (6, "농가 피드백 수렴 및 서비스 기능 개선",
     "C13 AI 영농비서 · 도움말·사용자 매뉴얼",
     [("C13 AI 영농비서",
       "농장 데이터 + 농진청 자료·병해충 DB(RAG 검색)를 결합한 생성형 AI 챗봇으로 "
       "관수·환경·수확·병해·데이터 입력 질의응답 및 피드백을 수렴한다(출처 배지 표기)."),
      ("도움말·사용자 매뉴얼",
       "영농 가이드·빠른 시작·무엇을 어디서·FAQ·용어·데이터 출처 배지 등 사용자 매뉴얼로 "
       "피드백을 반영해 사용성을 개선한다.")]),
    (7, "실증성과 분석 및 사업화 모델 도출",
     "C9 벤치마킹 · C10 투자 ROI",
     [("C9 벤치마킹",
       "RDA 표준·우수농가 기준 대비 비교, 개선 포인트, 월간 리포트 상세 연계로 실증성과를 분석한다."),
      ("C10 투자 ROI",
       "연매출(ERP)·연 에너지비 기반 ROI 산출, 농장 데이터 기반 투자안 비교·회수기간으로 "
       "사업화 모델과 확산 전략을 도출한다.")]),
]

doc = Document()
# 기본 폰트
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def set_kfont(run, size=10.5, bold=False, color=INK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)


# ── 제목 ──
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_kfont(p.add_run("KAASA Farmingsight"), 13, True, GREEN)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_kfont(p.add_run("7단계 실증 추진 로드맵 — 단계별 기능·화면 매핑"), 20, True, GREEN_DARK)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_kfont(p.add_run("스마트농업 경영최적화 의사결정 플랫폼  |  https://farmingsight.org"), 9.5, False, GRAY)
doc.add_paragraph()

# ── 요약 표 ──
p = doc.add_paragraph()
set_kfont(p.add_run("■ 단계 요약"), 12, True, GREEN_DARK)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for c, t in zip(hdr, ["단계", "추진 내용", "관련 화면"]):
    c.text = ""
    rp = c.paragraphs[0].add_run(t); set_kfont(rp, 10.5, True, RGBColor(0xFF,0xFF,0xFF))
    shade(c, "0F5132")
for n, title, screens, _ in STEPS:
    row = tbl.add_row().cells
    r0 = row[0].paragraphs[0].add_run(f"{n}단계"); set_kfont(r0, 10, True, GREEN_DARK)
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = row[1].paragraphs[0].add_run(title); set_kfont(r1, 10, False, INK)
    r2 = row[2].paragraphs[0].add_run(screens); set_kfont(r2, 9.5, False, GRAY)
# 열 너비
for row in tbl.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(3.5)
    row.cells[2].width = Inches(2.7)

doc.add_paragraph()

# ── 단계별 상세 ──
p = doc.add_paragraph()
set_kfont(p.add_run("■ 단계별 상세"), 12, True, GREEN_DARK)

for n, title, screens, items in STEPS:
    # 단계 제목 바
    p = doc.add_paragraph(); p.space_before = Pt(10)
    set_kfont(p.add_run(f"{n}단계  "), 13, True, GREEN)
    set_kfont(p.add_run(title), 13, True, GREEN_DARK)
    # 화면
    p = doc.add_paragraph()
    set_kfont(p.add_run("관련 화면  "), 9.5, True, GRAY)
    set_kfont(p.add_run(screens), 9.5, False, GREEN)
    # 본문(화면별)
    for label, body in items:
        bp = doc.add_paragraph(style="List Bullet")
        set_kfont(bp.add_run(f"[{label}]  "), 10.5, True, GREEN_DARK)
        set_kfont(bp.add_run(body), 10.5, False, INK)

OUT = os.path.join("out", "KAASA_7단계_실증추진_로드맵.docx")
os.makedirs("out", exist_ok=True)
doc.save(OUT)
print("저장 완료:", OUT)
